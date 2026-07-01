#!/usr/bin/env python3
"""Compile the project markdown into a single Word notebook (.docx).

Produces ``docs/ASRS_PdM_Notebook.docx`` — a stakeholder-ready book with a title page,
an explicit "Documents in this notebook" contents list, an auto-updating Word Table of
Contents field (right-click → Update Field, or it builds on open), and every source
markdown rendered in a logical reading order with headings, tables, code, lists,
blockquotes, and inline formatting preserved. Cross-references between docs are kept as
readable text.

Requires ``python-docx`` (a doc-build-only dependency): ``.venv/bin/pip install python-docx``.

Run:
    .venv/bin/python scripts/build_notebook.py
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent

# ---------------------------------------------------------------------------
# Notebook structure: (Part title, [(heading, relative path), ...])
# ---------------------------------------------------------------------------
STRUCTURE: List[Tuple[str, List[Tuple[str, str]]]] = [
    ("Part I — Overview & Foundations", [
        ("System Overview", "docs/SYSTEM_OVERVIEW.md"),
        ("Chapter 1 — Intro to the ASRS", "docs/notebook/01_intro_to_asrs.md"),
        ("PdM Methodology", "docs/notebook/methodology.md"),
    ]),
    ("Part II — Operating & Hosting", [
        ("Operator SOP", "docs/OPERATOR_SOP.md"),
        ("Hosting Resources", "docs/HOSTING_RESOURCES.md"),
        ("URL / Route Map", "docs/URL_MAP.md"),
    ]),
    ("Part III — Engineering & Reference", [
        ("Developer Guide", "docs/DEVELOPER_GUIDE.md"),
        ("Chapter 2 — Grafana Dashboards", "docs/notebook/02_grafana_dashboards.md"),
        ("Module ↔ Dashboard Mapping", "docs/mapping/module_dashboard_mapping.md"),
        ("Per-Module Health Methodology", "docs/MODULE_METHODOLOGY.md"),
        ("Audit & Hardening Report", "docs/AUDIT_REPORT.md"),
    ]),
    ("Part IV — Module Chapters", [
        ("Module 1 — Lift", "modules/lift/README.md"),
        ("Module 2 — Shuttle", "modules/shuttle/README.md"),
        ("Module 3 — Conveyor", "modules/conveyor/README.md"),
        ("Module 4 — Tracker / Position-Sensor", "modules/tracker/README.md"),
        ("Module 5 — Gate / Door-Actuator", "modules/gate/README.md"),
        ("Module 6 — Bin / Tote-Mechanical", "modules/bin_mech/README.md"),
        ("Module 7 — GTP Station + Scanner", "modules/gtp_station/README.md"),
        ("Module 8 — Decanting Station + Scanner", "modules/decant_station/README.md"),
        ("Module 9 — Network / Comms", "modules/network/README.md"),
        ("Module 10 — Controller / Compute", "modules/controller/README.md"),
        ("Module 11 — System-Wide Anomaly (Meta)", "modules/meta/README.md"),
    ]),
    ("Part V — Data Volume", [
        ("Chapter 3 — Data Volume", "docs/notebook/03_data_volume.md"),
    ]),
    ("Appendix", [
        ("Appendix A — Durable Conventions (CLAUDE.md)", "CLAUDE.md"),
    ]),
]

_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


# ---------------------------------------------------------------------------
# Inline + block rendering
# ---------------------------------------------------------------------------
def add_inline(paragraph, text: str) -> None:
    """Render **bold**, `code`, and [text](link) into runs on a paragraph."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        else:  # link
            lm = _LINK.match(tok)
            label, target = lm.group(1), lm.group(2)
            r = paragraph.add_run(label)
            r.font.color.rgb = RGBColor(0x1A, 0x4F, 0x8A)
            if target.startswith("http"):
                paragraph.add_run(f" ({target})").italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc, lines: List[str]) -> None:
    for ln in lines:
        p = doc.add_paragraph(style="No Spacing")
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
    doc.add_paragraph()  # spacer


def add_table(doc, rows: List[str]) -> None:
    def cells(line: str) -> List[str]:
        s = line.strip()
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        return [c.strip() for c in s.split("|")]

    header = cells(rows[0])
    body = [cells(r) for r in rows[2:]]  # rows[1] is the |---| separator
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:  # pragma: no cover - style name varies by template
        table.style = "Table Grid"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h.replace("**", ""))
        run.bold = True
    for r in body:
        cells_row = table.add_row().cells
        for i in range(ncol):
            val = r[i] if i < len(r) else ""
            cells_row[i].paragraphs[0].text = ""
            add_inline(cells_row[i].paragraphs[0], val)
    doc.add_paragraph()


def _is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", line)) and "-" in line


def render_markdown(doc, md: str) -> None:
    """Render a markdown string into the docx. Handles headings, lists, tables,
    fenced code, blockquotes, and inline formatting."""
    lines = md.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # fenced code
        if stripped.startswith("```"):
            block, i = [], i + 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i]); i += 1
            i += 1  # skip closing fence
            add_code_block(doc, block)
            continue

        # table (a pipe row followed by a separator row)
        if stripped.startswith("|") and i + 1 < n and _is_table_sep(lines[i + 1]):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            if len(block) >= 2:
                add_table(doc, block)
            continue

        # blank
        if not stripped:
            i += 1
            continue

        # horizontal rule -> spacer (part/file breaks are handled by the caller)
        if re.match(r"^\s*---+\s*$", line) or re.match(r"^\s*===+\s*$", line):
            i += 1
            continue

        # headings (file '#'->H2, '##'->H3, ... since Parts occupy H1)
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            level = len(hm.group(1))
            docx_level = min(level + 1, 6)  # shift so file title(#)->Heading2
            p = doc.add_heading("", level=docx_level)
            add_inline(p, re.sub(r"\s*#+\s*$", "", hm.group(2)))
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(re.sub(r"^\s*>\s?", "", lines[i])); i += 1
            p = doc.add_paragraph(style="Intense Quote" if _has_style(doc, "Intense Quote") else None)
            add_inline(p, " ".join(x for x in block if x.strip()))
            continue

        # bullet list
        bm = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if bm:
            indent = len(bm.group(1))
            style = "List Bullet 2" if indent >= 2 and _has_style(doc, "List Bullet 2") else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_inline(p, bm.group(2))
            i += 1
            continue

        # numbered list
        nm = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if nm:
            style = "List Number"
            p = doc.add_paragraph(style=style if _has_style(doc, style) else None)
            add_inline(p, nm.group(2))
            i += 1
            continue

        # normal paragraph (join wrapped lines until blank/structural)
        para = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", ">", "|", "```", "- ", "* "))
                    or re.match(r"^\d+\.\s", nxt) or re.match(r"^\s*---+\s*$", lines[i])):
                break
            para.append(nxt); i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(para))


def _has_style(doc, name: str) -> bool:
    try:
        _ = doc.styles[name]
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# TOC field + title page
# ---------------------------------------------------------------------------
def add_toc_field(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Table of Contents — right-click and choose “Update Field” to populate."
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_sep); r.append(placeholder); r.append(fld_end)


def build() -> Path:
    doc = Document()

    # Title page
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("ASRS Predictive Maintenance"); tr.bold = True; tr.font.size = Pt(28)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("System Notebook — Lenskart fulfilment plant, six-aisle ASRS")
    sr.font.size = Pt(14); sr.italic = True
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Compiled from the project markdown (operator, hosting, developer, "
                 "system, routes, methodology, module chapters, data volume).").italic = True
    doc.add_page_break()

    # Explicit contents list
    doc.add_heading("Documents in this notebook", level=1)
    doc.add_paragraph("This notebook compiles the source markdown in the order below. A "
                      "live, page-numbered Table of Contents follows (update it in Word).")
    for part, items in STRUCTURE:
        p = doc.add_paragraph(); p.add_run(part).bold = True
        for heading, rel in items:
            b = doc.add_paragraph(style="List Bullet")
            b.add_run(heading)
            b.add_run(f"  —  {rel}").italic = True
    doc.add_page_break()

    # Auto TOC field
    doc.add_heading("Table of Contents", level=1)
    add_toc_field(doc)
    doc.add_page_break()

    # Body
    missing: List[str] = []
    for part, items in STRUCTURE:
        doc.add_heading(part, level=1)
        for heading, rel in items:
            path = ROOT / rel
            doc.add_page_break()
            h = doc.add_heading("", level=2)
            add_inline(h, heading)
            src = doc.add_paragraph()
            src.add_run(f"Source: {rel}").italic = True
            if not path.exists():
                missing.append(rel)
                doc.add_paragraph(f"[missing: {rel}]")
                continue
            md = path.read_text(encoding="utf-8")
            # drop the file's own first H1 (we already printed a section heading)
            md = re.sub(r"\A\s*#\s+.*\n", "", md, count=1)
            render_markdown(doc, md)

    out = ROOT / "docs" / "ASRS_PdM_Notebook.docx"
    doc.save(out)
    if missing:
        print("WARNING: missing sources:", ", ".join(missing))
    return out


if __name__ == "__main__":
    out = build()
    print(f"Notebook written → {out} ({out.stat().st_size} bytes)")
