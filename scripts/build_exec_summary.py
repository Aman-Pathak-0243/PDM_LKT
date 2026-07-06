#!/usr/bin/env python3
"""Build the weekly Executive Summary (.docx) — a self-contained stakeholder report.

Renders charts + diagrams (matplotlib) from the live analytics data and assembles a
tightly-laid-out ~10-page Word document (python-docx). Reproducible each week: re-run
after refreshing database/analytics/ to update every figure and the fleet snapshot.

    python scripts/build_exec_summary.py
    -> docs/ASRS_PdM_Executive_Summary.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import pandas as pd

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------- #
# Palette (print-friendly on white)
# --------------------------------------------------------------------------- #
NAVY = "#14203a"; INK = "#1f2937"; MUTED = "#5a6878"
ACCENT = "#2563eb"; ACCENT_SOFT = "#e8f0fe"
OK = "#2ea043"; WATCH = "#d4a017"; WARN = "#e8770e"; CRIT = "#e5484d"
PANEL = "#eef2f7"; BORDER = "#c7d2e0"; ROWALT = "#f6f8fb"
TIER_COLOR = {"ok": OK, "watch": WATCH, "warn": WARN, "critical": CRIT, "unknown": MUTED}

FIGS = ROOT / "docs" / "_exec_figs"
FIGS.mkdir(parents=True, exist_ok=True)


def _hex(c):
    return RGBColor.from_string(c.lstrip("#"))


# --------------------------------------------------------------------------- #
# Live data (falls back to documented figures if analytics not yet built)
# --------------------------------------------------------------------------- #
def load_snapshot():
    f = ROOT / "database" / "analytics" / "component_health_timeseries.csv"
    fallback = {
        "components": 771, "avg": 90.0, "runs": 11,
        "tiers": {"ok": 603, "watch": 112, "warn": 34, "critical": 22},
        "per_module": [("gtp_station", 325, 90.9), ("network", 124, 91.9), ("shuttle", 124, 90.1),
                       ("tracker", 56, 85.9), ("gate", 52, 99.9), ("bin_mech", 41, 88.4),
                       ("decant_station", 19, 96.8), ("lift", 16, 63.7), ("meta", 7, 24.4),
                       ("conveyor", 6, 88.7), ("controller", 1, 100.0)],
        "top_risk": [], "live": False,
    }
    if not f.exists():
        return fallback
    df = pd.read_csv(f)
    df = df.sort_values("created_at").groupby(["module", "component_id"], as_index=False).tail(1)
    tiers = df.risk_tier.value_counts().to_dict()
    pm = (df.groupby("module").agg(n=("component_id", "size"), avg=("health_score", "mean"))
            .sort_values("n", ascending=False))
    per_module = [(m, int(r.n), round(float(r.avg), 1)) for m, r in pm.iterrows()]
    risk = df[df.risk_tier != "ok"].sort_values("health_score").head(8)
    top_risk = [(str(r.component_id), str(r.module), round(float(r.health_score), 1),
                 str(r.risk_tier), str(r.primary_cause) if pd.notna(r.primary_cause) else "")
                for _, r in risk.iterrows()]
    return {"components": int(len(df)), "avg": round(float(df.health_score.mean()), 1),
            "runs": int(df.run_uid.nunique()), "tiers": tiers, "per_module": per_module,
            "top_risk": top_risk, "live": True}


MODULE_TITLE = {
    "lift": "Lift", "shuttle": "Shuttle", "conveyor": "Conveyor", "tracker": "Tracker / Position-Sensor",
    "gate": "Gate / Door-Actuator", "bin_mech": "Bin / Tote-Mechanical", "gtp_station": "GTP Station + Scanner",
    "decant_station": "Decant Station + Scanner", "network": "Network / Comms", "controller": "Controller / Compute",
    "meta": "System-Wide Anomaly (Meta)",
}


# --------------------------------------------------------------------------- #
# Figures
# --------------------------------------------------------------------------- #
def _box(ax, x, y, w, h, text, fc, ec=BORDER, tc=INK, fs=9, bold=True):
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=2.2",
                                linewidth=1.2, edgecolor=ec, facecolor=fc, zorder=2))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs,
            color=tc, weight="bold" if bold else "normal", zorder=3, linespacing=1.25)


def _arrow(ax, x1, y1, x2, y2, color=ACCENT, ls="-", lw=1.6, text=None, tcol=MUTED):
    ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=12,
                                 linewidth=lw, color=color, linestyle=ls, zorder=1,
                                 shrinkA=1, shrinkB=1))
    if text:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 1.5, text, ha="center", va="bottom",
                fontsize=7.5, color=tcol, style="italic", zorder=3)


def fig_architecture(path):
    fig, ax = plt.subplots(figsize=(9.6, 4.5))
    ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis("off")
    # top pipeline
    xs = [1, 20.5, 40, 59.5, 79]; w = 17.5; y = 66; h = 20
    labels = ["Grafana\nDashboards\n(ops + error data)", "Fetch\nPlaywright + httpx",
              "Feature\nEngineering\nrates · z-scores", "Health Scoring\npenalty → 0–100\n+ risk tier",
              "RCA +\nCross-module\nflags"]
    fcs = ["#fdeee0", ACCENT_SOFT, ACCENT_SOFT, ACCENT_SOFT, ACCENT_SOFT]
    for i, (x, lab, fc) in enumerate(zip(xs, labels, fcs)):
        _box(ax, x, y, w, h, lab, fc, fs=8.3)
        if i:
            _arrow(ax, xs[i - 1] + w, y + h / 2, x, y + h / 2)
    # database (middle, wide)
    _box(ax, 20, 34, 60, 15, "CSV Database  ·  database/\nstore (live tables) · raw (gz snapshots) · analytics (EDA/ML) · archive · exports",
         PANEL, ec=ACCENT, tc=NAVY, fs=8.6)
    _arrow(ax, xs[4] + w / 2, y, 62, 49, text="persist verdicts + raw", tcol=MUTED)
    # consumers (bottom)
    _box(ax, 6, 6, 34, 16, "Web Dashboard\nFastAPI + Jinja2 · offline SVG charts\nOverview · per-module · storage · logs", "#eaf7ee", ec=OK, tc=NAVY, fs=8.3)
    _box(ax, 60, 6, 34, 16, "Automation\nAPScheduler (in-process)\nscheduled PdM runs", "#fdeee0", ec=WARN, tc=NAVY, fs=8.3)
    _arrow(ax, 30, 34, 23, 22, color=OK, text="reads", tcol=OK)
    _arrow(ax, 77, 22, xs[1] + w / 2, y, color=WARN, ls=(0, (4, 2)), text="triggers", tcol=WARN)
    _arrow(ax, 23, 22, xs[0] + w / 2, y, color=OK, ls=(0, (4, 2)), text="manual run", tcol=OK)
    fig.tight_layout(pad=0.3)
    fig.savefig(path, dpi=200, bbox_inches="tight"); plt.close(fig)


def fig_workflow(path):
    fig, ax = plt.subplots(figsize=(9.6, 2.5))
    ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis("off")
    steps = ["Trigger\n(manual / auto)", "Fetch\npanels", "Compute\nfeatures",
             "Score health\n0–100 → tier", "RCA +\ncross-flags", "Persist\n+ raw snapshot"]
    n = len(steps); w = 14.2; gap = (100 - n * w) / (n - 1); y = 52; h = 34
    for i, s in enumerate(steps):
        x = i * (w + gap)
        _box(ax, x, y, w, h, s, ACCENT_SOFT, fs=8.2)
        if i:
            _arrow(ax, (i - 1) * (w + gap) + w, y + h / 2, x, y + h / 2)
    ax.text(50, 20, "Cold-start → coarse tier bands (low confidence)   ·   Trend → regression RUL as the store accumulates history",
            ha="center", va="center", fontsize=8.2, color=NAVY, style="italic",
            bbox=dict(boxstyle="round,pad=0.5", fc="#fff8e8", ec=WATCH, lw=1))
    fig.tight_layout(pad=0.3)
    fig.savefig(path, dpi=200, bbox_inches="tight"); plt.close(fig)


def fig_donut(path, snap):
    order = ["ok", "watch", "warn", "critical"]
    vals = [snap["tiers"].get(t, 0) for t in order]
    cols = [TIER_COLOR[t] for t in order]
    fig, ax = plt.subplots(figsize=(4.3, 3.5))
    wedges, _ = ax.pie(vals, colors=cols, startangle=90, counterclock=False,
                       wedgeprops=dict(width=0.42, edgecolor="white", linewidth=2))
    ax.text(0, 0.12, f"{snap['components']:,}", ha="center", va="center", fontsize=23, weight="bold", color=NAVY)
    ax.text(0, -0.22, "components", ha="center", va="center", fontsize=10, color=MUTED)
    ax.legend(wedges, [f"{t.title()}  {snap['tiers'].get(t,0)}" for t in order],
              loc="center", bbox_to_anchor=(0.5, -0.06), ncol=2, frameon=False, fontsize=9,
              handlelength=1.1, columnspacing=1.4)
    ax.set(aspect="equal"); fig.tight_layout(pad=0.2)
    fig.savefig(path, dpi=200, bbox_inches="tight"); plt.close(fig)


def fig_module_bar(path, snap):
    data = sorted(snap["per_module"], key=lambda r: r[1])
    names = [MODULE_TITLE.get(m, m) for m, _, _ in data]
    counts = [n for _, n, _ in data]
    avgs = [a for _, _, a in data]
    def tier(a): return "ok" if a >= 85 else "watch" if a >= 65 else "warn" if a >= 40 else "critical"
    cols = [TIER_COLOR[tier(a)] for a in avgs]
    fig, ax = plt.subplots(figsize=(9.4, 3.9))
    bars = ax.barh(names, counts, color=cols, edgecolor="white", height=0.68)
    ax.set_xlabel("Components monitored", fontsize=9, color=MUTED)
    ax.tick_params(labelsize=9)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)
    ax.spines["left"].set_color(BORDER); ax.spines["bottom"].set_color(BORDER)
    mx = max(counts)
    for b, c, a in zip(bars, counts, avgs):
        ax.text(b.get_width() + mx * 0.012, b.get_y() + b.get_height() / 2,
                f"{c}   ·   avg {a:.0f}", va="center", ha="left", fontsize=8.2, color=INK)
    ax.set_xlim(0, mx * 1.22)
    fig.tight_layout(pad=0.4)
    fig.savefig(path, dpi=200, bbox_inches="tight"); plt.close(fig)


# --------------------------------------------------------------------------- #
# docx helpers
# --------------------------------------------------------------------------- #
def set_cell_bg(cell, hexcolor):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(sh)


def _no_wrap_pad(cell, top=40, bottom=40, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa"); m.append(e)
    tcPr.append(m)


def cell_para(cell, clear=True):
    if clear:
        cell.text = ""
    return cell.paragraphs[0]


def run(p, text, size=10.5, bold=False, color=INK, italic=False, font="Calibri"):
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.name = font; r.font.color.rgb = _hex(color); return r


def set_table_borders(table, color=BORDER, sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color.lstrip("#"))
        borders.append(e)
    tblPr.append(borders)


def add_rule(doc, color=ACCENT, size=18):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True  # bind the rule to the content that follows
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color.lstrip("#"))
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def h1(doc, text, num=None):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(13); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    if num:
        run(p, f"{num}  ", size=14, bold=True, color=ACCENT)
    run(p, text, size=14, bold=True, color=NAVY)
    add_rule(doc, ACCENT, 14)


def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    run(p, text, size=11, bold=True, color=ACCENT)
    return p


def body(doc, text, size=11, space_after=7, align=None, keep_next=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_with_next = keep_next
    if align:
        p.alignment = align
    run(p, text, size=size)
    return p


def bullet(doc, label, text, size=11):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.16
    if label:
        run(p, label + " ", size=size, bold=True, color=NAVY)
    run(p, text, size=size)
    return p


def module_block(doc, idx, title, ctype, count, md):
    """One compact per-module card: name, what it does, signals (features), formulas."""
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    run(p, f"{idx}.  {title}", size=10.5, bold=True, color=NAVY)
    run(p, f"      {ctype} · {count} units", size=8.6, color=MUTED)
    summ = (md.get("summary") or md.get("description") or "").strip()
    summ = summ.split(". ")[0].strip().rstrip(".")
    if len(summ) > 195:
        summ = summ[:193].rstrip() + "…"
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(1.5); sp.paragraph_format.line_spacing = 1.12
    sp.paragraph_format.keep_with_next = True
    run(sp, summ + ".", size=9.6)
    sigs = md.get("signals") or []
    if sigs:
        gp = doc.add_paragraph(); gp.paragraph_format.space_after = Pt(1.5); gp.paragraph_format.line_spacing = 1.14
        run(gp, "Signals (features):  ", size=9.3, bold=True, color=ACCENT)
        for i, s in enumerate(sigs):
            what = (s.get("what") or "").split(".")[0].strip()
            if len(what) > 46:
                what = what[:44].rstrip() + "…"
            run(gp, s.get("name", ""), size=9.3, bold=True, color=INK)
            run(gp, f" — {what}", size=9.3, color=MUTED)
            if i < len(sigs) - 1:
                run(gp, "   ·   ", size=9.3, color="#9fb0c4")
    fs = [f for f in (md.get("formulas") or [])
          if not str(f.get("name", "")).strip().lower().startswith("health")]
    if fs:
        fp = doc.add_paragraph(); fp.paragraph_format.space_after = Pt(2); fp.paragraph_format.line_spacing = 1.14
        run(fp, "Formulas:  ", size=9.1, bold=True, color=ACCENT)
        for i, f in enumerate(fs):
            run(fp, f"{f.get('name')} = {f.get('formula')}", size=8.7, font="Consolas", color=INK)
            if i < len(fs) - 1:
                run(fp, "       ", size=9)


def add_image(doc, path, width_in, caption=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_in))
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(7)
        run(c, caption, size=8.5, italic=True, color=MUTED)


def add_footer(section, title):
    p = section.footer.paragraphs[0]; p.text = ""
    run(p, title + "      ", size=8, color=MUTED)
    # page number field
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rpr.append(sz); r.append(rpr)
    t = OxmlElement("w:t"); t.text = "1"; r.append(t); fld.append(r)
    p._p.append(fld)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def build():
    snap = load_snapshot()
    fig_architecture(FIGS / "arch.png")
    fig_workflow(FIGS / "flow.png")
    fig_donut(FIGS / "donut.png", snap)
    fig_module_bar(FIGS / "modbar.png", snap)

    doc = Document()
    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.font.color.rgb = _hex(INK)
    sec = doc.sections[0]
    sec.page_height = Inches(11.69); sec.page_width = Inches(8.27)  # A4
    for m in ("top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.7))
    for m in ("left_margin", "right_margin"):
        setattr(sec, m, Inches(0.75))
    sec.footer_distance = Inches(0.4)
    add_footer(sec, "ASRS Predictive Maintenance  ·  Weekly Executive Summary  ·  Confidential")
    CW = 6.77  # content width in inches

    # ---- Cover band ----
    band = doc.add_table(rows=1, cols=1); band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.columns[0].width = Inches(CW)
    c = band.rows[0].cells[0]; set_cell_bg(c, NAVY); _no_wrap_pad(c, 220, 220, 220, 220)
    c.width = Inches(CW)
    p = cell_para(c); p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after = Pt(2)
    run(p, "ASRS Predictive Maintenance System", size=20, bold=True, color="#ffffff")
    p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(0)
    run(p2, "Weekly Progress Report — Executive Summary", size=12, color="#c7d6f5")
    p3 = c.add_paragraph(); p3.paragraph_format.space_before = Pt(6); p3.paragraph_format.space_after = Pt(0)
    run(p3, "Lenskart fulfilment plant · six-aisle ASRS      |      Week ending 6 July 2026", size=9.5, color="#aab8d6")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- Abstract ----
    body(doc,
         "This report summarises the ASRS Predictive Maintenance (PdM) system and this week's "
         "progress. The system infers the health of every physical component in the automated "
         "storage & retrieval system purely from Grafana operational and error data — there is no "
         "maintenance logbook — and predicts what is degrading, how urgently, and why, so "
         "maintenance is scheduled proactively instead of reacting to breakdowns. The module set is "
         "complete (11/11) and the system is running against live plant data. This summary covers "
         "the system's purpose, architecture, and methodology; the eleven equipment modules; the "
         "current live fleet state; this week's deliverables; and the outlook.",
         size=11, space_after=5)

    # ---- KPI tiles ----
    kpis = [("11", "Modules (complete)"), (f"{snap['components']:,}", "Components monitored"),
            (f"{snap['avg']:.0f}", "Fleet health / 100"), (f"{snap['tiers'].get('critical',0)}", "Critical now"),
            ("36", "Automated tests"), ("CSV", "Single database store")]
    kt = doc.add_table(rows=1, cols=len(kpis)); kt.alignment = WD_TABLE_ALIGNMENT.CENTER
    kt.autofit = False
    for i, (val, lab) in enumerate(kpis):
        cell = kt.rows[0].cells[i]; cell.width = Inches(CW / len(kpis))
        set_cell_bg(cell, ACCENT_SOFT if i % 2 == 0 else PANEL); _no_wrap_pad(cell, 70, 70, 60, 60)
        pv = cell_para(cell); pv.alignment = WD_ALIGN_PARAGRAPH.CENTER; pv.paragraph_format.space_after = Pt(0)
        run(pv, val, size=17, bold=True, color=ACCENT if i != 3 else CRIT)
        pl = cell.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER; pl.paragraph_format.space_after = Pt(0)
        run(pl, lab, size=7.6, color=MUTED)
    set_table_borders(kt, "#ffffff", 6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- 1. Overview ----
    h1(doc, "The Problem & Our Approach", "1")
    body(doc,
         "The plant runs a six-aisle ASRS — lifts, shuttles, conveyors, gates, bins, scanners, "
         "decant stations, a comms layer and a controller. When a component degrades the usual "
         "outcomes are an unplanned breakdown (lost throughput, missed dispatch) or over-servicing "
         "of parts that were fine. Crucially there is no maintenance logbook or failure history to "
         "train a classical model on — but the plant already produces rich Grafana operational and "
         "error data. This system turns that data into predictive maintenance using four ideas:")
    bullet(doc, "Condition monitoring, not failure regression —",
           "each component starts at 100 health and loses weighted, capped penalties for each "
           "unhealthy signal (error/fault rate, severity mix, recurrence, peer & self deviation): "
           "health = clamp(100 − Σ penalties, 0, 100), mapped to a tier (ok / watch / warn / critical).")
    bullet(doc, "The store beats the 2-day window —",
           "Grafana retains ~2 days, so every run snapshots each component's metrics into a "
           "longitudinal store. Over runs this becomes a history far longer than any single fetch, "
           "enabling recurrence, persistence and trend-based remaining-useful-life (RUL).")
    bullet(doc, "Per-component → per-module → system —",
           "each physical unit gets its own verdict; a module tile shows its worst component; and a "
           "Meta layer collapses correlated failures across modules into one compound incident with a "
           "likely common cause, so operators chase one root issue rather than ten symptoms.")
    bullet(doc, "Explainable & auditable —",
           "every flag carries a root-cause attribution (dominant signals + error mix) and a "
           "confidence, and every run is a traceable trigger, so a recommendation can always be "
           "justified to the maintenance team.")

    # ---- 2. Architecture ----
    h1(doc, "System Architecture", "2")
    body(doc,
         "A single process on one LAN PC serves the dashboard and runs automation. A PdM run is a "
         "four-stage pipeline — fetch → features → health → persist — and the CSV database is the "
         "shared source of truth: written by every run, read by the dashboard.", space_after=3,
         keep_next=True)
    add_image(doc, FIGS / "arch.png", CW,
              "Figure 1 — End-to-end architecture: Grafana data is fetched, scored, explained and persisted to the CSV database; "
              "the dashboard reads it and automation (or an operator) triggers runs.")
    body(doc,
         "Stack: Python 3.11+, FastAPI + Uvicorn (web), APScheduler (in-process automation), "
         "Playwright (Chromium) + httpx (Grafana fetch), numpy / pandas / scikit-learn (modelling), "
         "server-rendered Jinja2 with vendored SVG charts (fully offline on the LAN — no CDN). A "
         "plugin registry lets each equipment type self-register as a module, so adding equipment "
         "needs no changes to the core.", space_after=3)

    # ---- 3. Methodology & workflow ----
    h1(doc, "How a Run Works — Methodology", "3")
    add_image(doc, FIGS / "flow.png", CW,
              "Figure 2 — The PdM run pipeline. Predictions are labelled cold-start or trend and carry a confidence that rises as history grows.")
    body(doc,
         "Every trigger (manual or scheduled) fetches each mapped panel, derives normalised, "
         "dimensionless features (rates, ratios, robust z-scores vs the unit's own baseline and its "
         "peers — never raw counts, so a busy unit isn't punished for being busy), scores health, "
         "attributes a root cause, and persists the run. Two regimes are always labelled: "
         "cold-start (little history → coarse, low-confidence tier bands) and trend (enough history "
         "→ a fitted health trajectory projects when a maintenance threshold is crossed). Running it "
         "regularly is what makes it predictive — each run sharpens the forecast.", space_after=4)

    # ---- 4. What it monitors ----
    h1(doc, "What It Monitors — 11 Modules", "4")
    signals = {
        "gtp_station": "scanner misread rate + station pick-discrepancy rate",
        "network": "network downtime % + today-vs-window spike + aisle clustering",
        "shuttle": "errors per million cycles + severity + cycles-based RUL",
        "tracker": "mislocated-tote clustering + cross-run recurrence",
        "gate": "stuck-non-closed latency + cross-run persistence",
        "bin_mech": "block-age + chronic-slot history + recurrence",
        "decant_station": "scanner misread rate + station status/throughput",
        "lift": "error rate + severity + mechanical-wear mix + peer deviation",
        "meta": "cross-module compound-risk (co-occurrence + causal chains)",
        "conveyor": "congestion (queue vs limit) + stall detection",
        "controller": "CPU utilisation % + sustained-high + trend",
    }
    tbl = doc.add_table(rows=1, cols=4); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, BORDER, 4)
    widths = [0.4, 2.35, 3.1, 0.9]
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(["#", "Module", "Leading health signal", "Units"]):
        hdr[i].width = Inches(widths[i]); set_cell_bg(hdr[i], NAVY); _no_wrap_pad(hdr[i])
        pp = cell_para(hdr[i]); pp.paragraph_format.space_after = Pt(0)
        run(pp, htext, size=9, bold=True, color="#ffffff")
    counts = {m: n for m, n, _ in snap["per_module"]}
    for idx, (m, _, _) in enumerate(snap["per_module"], 1):
        cells = tbl.add_row().cells
        vals = [str(idx), MODULE_TITLE.get(m, m), signals.get(m, ""), str(counts.get(m, ""))]
        for i, v in enumerate(vals):
            cells[i].width = Inches(widths[i]); _no_wrap_pad(cells[i])
            if idx % 2 == 0:
                set_cell_bg(cells[i], ROWALT)
            pp = cell_para(cells[i]); pp.paragraph_format.space_after = Pt(0)
            run(pp, v, size=8.7, bold=(i == 1), color=NAVY if i == 1 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    add_image(doc, FIGS / "modbar.png", CW,
              "Figure 3 — Components monitored per module (bar length) and average health (label & colour: green ok · amber watch · orange warn · red critical).")

    # ---- 5. Per-module methodology (pulled live from each module's methodology dict) ----
    h1(doc, "Per-Module Methodology, Features & Formulas", "5")
    body(doc,
         "All eleven modules share one scoring model — health = clamp(100 − Σ capped penalties, 0, 100), "
         "mapped to a risk tier, with each prediction labelled cold-start or trend. What differs per module "
         "is the equipment it watches and the signals (features) it derives from the Grafana data. Each block "
         "below states what the module does, the features it uses, and its distinctive formulas (the shared "
         "health formula above is not repeated).", space_after=6)
    try:
        import modules  # noqa: F401  self-register
        from core.registry import all_modules, module_methodology
        for i, m in enumerate(all_modules(), 1):
            module_block(doc, i, MODULE_TITLE.get(m.name, m.title), m.component_type,
                         counts.get(m.name, "—"), module_methodology(m))
    except Exception as exc:  # pragma: no cover - environment without .env
        body(doc, f"(Per-module methodology unavailable in this environment: {exc})", size=9, color=MUTED)

    # ---- 6. Current state ----
    h1(doc, "Current System State (live data)", "6")
    body(doc,
         "The snapshot below is computed from the most recent scored run of the live fleet. The "
         "donut shows how the monitored components split across the four risk tiers; the table "
         "lists the specific units with the lowest health for the maintenance team to review first. "
         "These figures update automatically every time a PdM run completes.", space_after=6)
    st = doc.add_table(rows=1, cols=2); st.alignment = WD_TABLE_ALIGNMENT.CENTER
    st.columns[0].width = Inches(2.6); st.columns[1].width = Inches(CW - 2.6)
    lc = st.rows[0].cells[0]; lc.width = Inches(2.6)
    add_image_into_cell(lc, FIGS / "donut.png", 2.5)
    rc = st.rows[0].cells[1]; rc.width = Inches(CW - 2.6)
    rp = cell_para(rc); rp.paragraph_format.space_after = Pt(3)
    run(rp, f"Across {snap['components']:,} monitored components the fleet averages "
            f"{snap['avg']:.0f}/100 health. ", size=10, )
    run(rp, f"{snap['tiers'].get('critical',0)} critical and {snap['tiers'].get('warn',0)} warn "
            "components are surfaced for attention; the rest sit in watch/ok.", size=10)
    if snap["top_risk"]:
        rp2 = rc.add_paragraph(); rp2.paragraph_format.space_after = Pt(2)
        run(rp2, "Top at-risk components", size=9.5, bold=True, color=NAVY)
        rt = rc.add_table(rows=1, cols=3); set_table_borders(rt, BORDER, 3)
        for i, htext in enumerate(["Component", "Module", "Health"]):
            hc = rt.rows[0].cells[i]; set_cell_bg(hc, PANEL); _no_wrap_pad(hc, 20, 20, 60, 60)
            pp = cell_para(hc); pp.paragraph_format.space_after = Pt(0)
            run(pp, htext, size=8, bold=True, color=NAVY)
        for cid, mod, hs, tier, _cause in snap["top_risk"][:5]:
            cells = rt.add_row().cells
            for i, v in enumerate([cid, MODULE_TITLE.get(mod, mod).split(" /")[0], f"{hs:.0f}"]):
                _no_wrap_pad(cells[i], 18, 18, 60, 60)
                pp = cell_para(cells[i]); pp.paragraph_format.space_after = Pt(0)
                col = TIER_COLOR.get(tier, INK) if i == 2 else INK
                run(pp, v, size=8, bold=(i == 2), color=col)
    body(doc,
         "Note: all components are currently in the cold-start regime — the store is early in its "
         "accumulation. As scheduled runs continue, confidence and trend-based RUL sharpen "
         "automatically; the Meta module already flags the aisles carrying compound risk.",
         size=9.5, space_after=4)

    # ---- 7. This week's progress ----
    h1(doc, "This Week's Progress", "7")
    deliverables = [
        ("Graphical Overview analytics tab", "Delivered",
         "New tab on the Overview page: KPI row + 7 dependency-free SVG charts (fleet trend, status donut, per-module risk, score distribution, aisle×module heatmap, top at-risk, time-to-maintenance)."),
        ("Single CSV 'database' store", "Delivered",
         "Consolidated all persistence into one database/ folder (store · analytics · raw · archive · exports); migrated live history into it."),
        ("Raw data capture per run", "Delivered",
         "Every run now snapshots the raw fetched Grafana data (gzipped CSV per panel) — full audit trail; nothing is discarded."),
        ("Analytics / EDA-ML datasets", "Delivered",
         "Builder emits tidy time-series + per-module feature matrices + data dictionary for future trend analysis, EDA and ML."),
        ("Dockerisation + deploy guide", "Delivered",
         "One-container image (dashboard + automation) with host-mounted data; step-by-step LAN deployment guide added to the README."),
        ("Correctness hardening", "Delivered",
         "Adversarial multi-agent code review; fixed 3 confirmed defects (window-unit parsing, TTM boundary, module sort) and a missing PyYAML dependency that would break fresh installs."),
        ("Documentation & tests", "Delivered",
         "Refreshed all docs + rebuilt the Word notebook; automated test suite expanded to 36 (all passing)."),
    ]
    dt = doc.add_table(rows=1, cols=3); dt.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(dt, BORDER, 4)
    dwidths = [2.15, 0.85, CW - 3.0]
    for i, htext in enumerate(["Deliverable", "Status", "What it means"]):
        hc = dt.rows[0].cells[i]; hc.width = Inches(dwidths[i]); set_cell_bg(hc, NAVY); _no_wrap_pad(hc)
        pp = cell_para(hc); pp.paragraph_format.space_after = Pt(0)
        run(pp, htext, size=9, bold=True, color="#ffffff")
    for idx, (name, status, desc) in enumerate(deliverables):
        cells = dt.add_row().cells
        for i, v in enumerate([name, status, desc]):
            cells[i].width = Inches(dwidths[i]); _no_wrap_pad(cells[i])
            if idx % 2 == 0:
                set_cell_bg(cells[i], ROWALT)
            pp = cell_para(cells[i]); pp.paragraph_format.space_after = Pt(0)
            if i == 1:
                set_cell_bg(cells[i], "#e6f5ea")
                run(pp, "✓ " + v, size=8.6, bold=True, color=OK)
            else:
                run(pp, v, size=8.7, bold=(i == 0), color=NAVY if i == 0 else INK)

    # ---- 8. Data & deployment ----
    h1(doc, "Data Store & Deployment", "8")
    h2(doc, "CSV-only, analysis-ready")
    body(doc,
         "Storage is CSV-only (MySQL designed but dormant behind a permission gate). Everything "
         "lives in one database/ folder: store/ (the live longitudinal tables — one row per "
         "component per run), raw/ (per-run gzipped snapshots of the exact fetched data), "
         "analytics/ (flat, tidy time-series + per-module feature matrices ready for pandas / "
         "scikit-learn), plus archive/ and exports/. A data dictionary ships with it. All feature "
         "and root-cause data is kept as JSON so future ML/analytics never needs a schema migration.",
         space_after=4)
    layout = [
        ("store/", "Live longitudinal tables — one row per component per run (the health history)."),
        ("raw/", "Per-run gzipped snapshots of the exact fetched Grafana data (full audit trail)."),
        ("analytics/", "Flat, tidy time-series + per-module feature matrices — ready for EDA / ML."),
        ("archive/ · exports/", "Aged-out rows and one-off CSV / JSON / Excel exports."),
    ]
    lt = doc.add_table(rows=1, cols=2); lt.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(lt, BORDER, 4)
    lw = [1.7, CW - 1.7]
    for i, htext in enumerate(["database/ folder", "Contents"]):
        hc = lt.rows[0].cells[i]; hc.width = Inches(lw[i]); set_cell_bg(hc, NAVY); _no_wrap_pad(hc)
        pp = cell_para(hc); pp.paragraph_format.space_after = Pt(0)
        run(pp, htext, size=9, bold=True, color="#ffffff")
    for idx, (folder, desc) in enumerate(layout):
        cells = lt.add_row().cells
        for i, v in enumerate([folder, desc]):
            cells[i].width = Inches(lw[i]); _no_wrap_pad(cells[i])
            if idx % 2 == 0:
                set_cell_bg(cells[i], ROWALT)
            pp = cell_para(cells[i]); pp.paragraph_format.space_after = Pt(0)
            run(pp, v, size=8.9, bold=(i == 0), color=NAVY if i == 0 else INK,
                font="Consolas" if i == 0 else "Calibri")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    h2(doc, "One-container delivery")
    body(doc,
         "The system is packaged as a single Docker container (dashboard + automation) with the "
         "database and logs bind-mounted to the host for easy backup. It binds the LAN and is "
         "reached at http://<host-ip>:8800; automation runs inside the container independent of any "
         "browser. A step-by-step deployment guide (build, run, firewall, operations) is in the "
         "README for installation on the target plant PC.", space_after=4)

    # ---- 9. Roadmap ----
    h1(doc, "Outlook & Next Steps", "9")
    bullet(doc, "Accumulate history —",
           "keep scheduled runs going so components move from cold-start to trend, unlocking sharper "
           "RUL and higher-confidence predictions.")
    bullet(doc, "MySQL when provisioned —",
           "flip one flag and run a one-command migration; no application code changes (the storage "
           "layer is already abstracted).")
    bullet(doc, "Operator feedback loop —",
           "optional 'mark maintenance done' acknowledgements to annotate flags and reset baselines "
           "(never drives detection).")
    bullet(doc, "Leverage the raw/analytics datasets —",
           "the newly-captured raw data and EDA/ML extracts open the door to learned failure models "
           "as labelled outcomes accrue.")

    out = ROOT / "docs" / "ASRS_PdM_Executive_Summary.docx"
    doc.save(out)
    return out, snap


def add_image_into_cell(cell, path, width_in):
    p = cell_para(cell); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width_in))


if __name__ == "__main__":
    out, snap = build()
    print(f"Executive summary written -> {out} ({out.stat().st_size} bytes)")
    print(f"  data: {'LIVE' if snap['live'] else 'FALLBACK'} · {snap['components']} components · avg {snap['avg']}")
